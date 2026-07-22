# -*- coding: utf-8 -*-

import os
import sys
import docx
import pandas as pd

# This script is designed to search through Word documents in a specified input directory.
# The script will look for specific keywords or phrases and output the match times in each document to a specified output directory.

print_detailed_info = False  # Set to True to print detailed information during processing

# get the input directory and output directory from command line arguments
if len(sys.argv) != 5:
    print("Usage: python search_keywords.py <file_type> <input_directory> <output_directory> <keywords>")
    sys.exit(1)

file_type = sys.argv[1]
input_directory = sys.argv[2]
output_directory = sys.argv[3]
keywords = sys.argv[4]

if file_type not in ['docx', 'xlsx']:
    print(f"Error: Unsupported file type - {file_type}. Supported types are 'docx' and 'xlsx'.")
    sys.exit(1)

# get keywords from the keywords file
try:
    with open(keywords, 'r', encoding='utf-8') as f:
        keyword_list = [line.strip() for line in f if line.strip()]
        print(f"\nKeywords loaded from {keywords}:\n {keyword_list}")
except FileNotFoundError:
    print(f"Error: File not found - {keywords}")
    sys.exit(1)

# prepare a list to hold the results, each result will be a tuple of (input_directory, filename, keyword, times_found)
raw_results = []
failed_files = []

# clean output files
if file_type == 'docx':
    output_file_path = os.path.join(output_directory, 'word_search_results.xlsx')
    output_failed_file_path = os.path.join(output_directory, 'word_failed_files.txt')
elif file_type == 'xlsx':
    output_file_path = os.path.join(output_directory, 'excel_search_results.xlsx')
    output_failed_file_path = os.path.join(output_directory, 'excel_failed_files.txt')
print("\nCleaning up existing output files if they exist...")
if os.path.exists(output_file_path):
    os.remove(output_file_path)
    print(f"Existing output file removed: {output_file_path}")
if os.path.exists(output_failed_file_path):
    os.remove(output_failed_file_path)
    print(f"Existing failed files log removed: {output_failed_file_path}")

# check if there are multiple input directories
if os.path.isfile(input_directory):
    # read the input directory from the file
    with open(input_directory, 'r', encoding='utf-8') as f:
        input_directories = f.read().strip().splitlines()
    print(f"\nInput directories loaded from {input_directory}:\n {input_directories}")
else:
    # if the input directory is not a file, check if it is a valid directory
    if not os.path.isdir(input_directory):
        print(f"Error: Input directory does not exist - {input_directory}")
        sys.exit(1)
    else:
        input_directories = [input_directory]
        print(f"\nInput directory: {input_directory}")

# to get text in table format
def get_table_text(doc):
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_texts.append(cell.text)
    return table_texts

# to get header and footer text
def get_header_footer_content(doc):
    headers_footers = []
    for _idx, section in enumerate(doc.sections):
        header = section.header
        footer = section.footer
        for p in header.paragraphs:
            headers_footers.append(p.text.strip())
            # print(f"Header text: {p.text.strip()}")
        for p in footer.paragraphs:
            headers_footers.append(p.text.strip())
            # print(f"Footer text: {p.text.strip()}")
    return headers_footers

# search through Word documents in the input directory
for input_directory in input_directories:
    for filename in os.listdir(input_directory):
        if file_type == 'docx' and (filename.endswith('.docx') or filename.endswith('.doc')):
            file_path = os.path.join(input_directory, filename)
            # Here you would add the logic to open the Word document and search for keywords
            # For example, using python-docx library to read the document and search for keywords
            # This is a placeholder for the actual search logic
            try:
                print(f"Searching in {file_path}...")
                doc = docx.Document(file_path)
            except Exception as e:
                print("") # print a blank line for better readability
                print(f"[WARNING] Failed to process {file_path}: {e}")
                print("")
                failed_files.append(file_path)
                doc = None
            
            if doc is not None:
                for keyword in keyword_list:
                    times_found_total = 0
                    for para in doc.paragraphs:
                        times_found_in_para = para.text.lower().count(keyword.lower())
                        times_found_total += times_found_in_para
                    for table_text in get_table_text(doc):
                        times_found_in_table = table_text.lower().count(keyword.lower())
                        times_found_total += times_found_in_table
                    for header_footer_text in get_header_footer_content(doc):
                        times_found_in_header_footer = header_footer_text.lower().count(keyword.lower())
                        times_found_total += times_found_in_header_footer
                    if times_found_total > 0 and print_detailed_info:
                        print(f"Found '{keyword}' {times_found_total} times in {file_path}")
                    raw_results.append((input_directory, filename, keyword, times_found_total))

        elif file_type == 'xlsx' and (filename.endswith('.xlsx') or filename.endswith('.xls')):
            file_path = os.path.join(input_directory, filename)
            # Here you would add the logic to open the Excel document and search for keywords
            # This is a placeholder for the actual search logic
            try:
                print(f"Searching in {file_path}...")
                xls = pd.ExcelFile(file_path)
            except Exception as e:
                print("") # print a blank line for better readability
                print(f"[WARNING] Failed to process {file_path}: {e}")
                print("")
                failed_files.append(file_path)
                xls = None
            
            if xls is not None:
                for sheet_name in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet_name)
                    for keyword in keyword_list:
                        times_found_total = df.apply(lambda row: row.astype(str).str.contains(keyword, case=False).sum(), axis=1).sum()
                        if times_found_total > 0 and print_detailed_info:
                            print(f"Found '{keyword}' {times_found_total} times in {file_path} (Sheet: {sheet_name})")
                        raw_results.append((input_directory, filename, keyword, times_found_total))

# convert the raw results to a table format:
# each row index will be filename, and each column index will be a keyword, 
# with the cell value being the number of times the keyword was found in that file
columns = ['Input Directory', 'Filename'] + keyword_list
results_table = pd.DataFrame(columns=columns)
for input_directory, filename, keyword, times_found in raw_results:
    if filename not in results_table['Filename'].values:
        results_table = pd.concat([results_table, pd.DataFrame([[input_directory, filename] + [0]*len(keyword_list)], columns=columns)], ignore_index=True)
    results_table.loc[results_table['Filename'] == filename, keyword] = times_found

# highlight the cells (set the cell color to yellow) that have values greater than 0
def highlight_non_zero(val):
   return f'background-color: {"yellow"}' if (isinstance(val, int) and val > 0) else None
results_table = results_table.style.applymap(highlight_non_zero)

# export the results to a excel file in the output directory
results_table.to_excel(output_file_path, index=False)
print(f"\nSearch results exported to {output_file_path}")

if failed_files:
    print(f"\n[WARNING] Failed to process the following files:")
    for file_path in failed_files:
        print(f" - {file_path}")

    with open(output_failed_file_path, 'w', encoding='utf-8') as f:
        for file_path in failed_files:
            f.write(f"{file_path}\n")